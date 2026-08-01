from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
from pathlib import Path

OUT = Path(r"C:\Users\redcl\ai project\선거캠프OS_통합운영플랫폼_기획안.docx")
NAVY = "102A43"; BLUE = "2563EB"; CYAN = "0891B2"; PALE = "EAF2FF"
LIGHT = "F3F6FA"; MID = "64748B"; DARK = "172033"; WHITE = "FFFFFF"

doc = Document()
sec = doc.sections[0]
sec.page_width, sec.page_height = Inches(8.5), Inches(11)
sec.top_margin = sec.bottom_margin = Inches(0.72)
sec.left_margin = sec.right_margin = Inches(0.85)
sec.header_distance = sec.footer_distance = Inches(0.35)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Malgun Gothic"; normal.font.size = Pt(10.5); normal.font.color.rgb = RGBColor.from_string(DARK)
normal._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
normal.paragraph_format.space_after = Pt(7); normal.paragraph_format.line_spacing = 1.25
for name, size, color, before, after in [
    ("Heading 1", 18, NAVY, 18, 8), ("Heading 2", 14, BLUE, 14, 6), ("Heading 3", 11.5, NAVY, 10, 4)
]:
    st = styles[name]; st.font.name = "Malgun Gothic"; st.font.size = Pt(size); st.font.bold = True
    st.font.color.rgb = RGBColor.from_string(color); st._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    st.paragraph_format.space_before = Pt(before); st.paragraph_format.space_after = Pt(after); st.paragraph_format.keep_with_next = True

for style_name in ["Proposal Lead", "Callout"]:
    if style_name not in styles:
        styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
styles["Proposal Lead"].font.name = "Malgun Gothic"; styles["Proposal Lead"].font.size = Pt(13)
styles["Proposal Lead"].font.bold = True; styles["Proposal Lead"].font.color.rgb = RGBColor.from_string(NAVY)
styles["Proposal Lead"].paragraph_format.space_after = Pt(12); styles["Proposal Lead"].paragraph_format.line_spacing = 1.3
styles["Callout"].font.name = "Malgun Gothic"; styles["Callout"].font.size = Pt(10.5)
styles["Callout"].font.bold = True; styles["Callout"].font.color.rgb = RGBColor.from_string(BLUE)
styles["Callout"].paragraph_format.space_before = Pt(8); styles["Callout"].paragraph_format.space_after = Pt(10)

def font(run, size=None, bold=None, color=None):
    run.font.name = "Malgun Gothic"; run._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    if size: run.font.size = Pt(size)
    if bold is not None: run.bold = bold
    if color: run.font.color.rgb = RGBColor.from_string(color)

def shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr(); shd = tcPr.find(qn("w:shd"))
    if shd is None: shd = OxmlElement("w:shd"); tcPr.append(shd)
    shd.set(qn("w:fill"), fill)

def margins(cell, top=100, start=120, bottom=100, end=120):
    tc = cell._tc.get_or_add_tcPr(); tcMar = tc.first_child_found_in("w:tcMar")
    if tcMar is None: tcMar = OxmlElement("w:tcMar"); tc.append(tcMar)
    for side, val in [("top",top),("start",start),("bottom",bottom),("end",end)]:
        node = tcMar.find(qn(f"w:{side}"))
        if node is None: node = OxmlElement(f"w:{side}"); tcMar.append(node)
        node.set(qn("w:w"), str(val)); node.set(qn("w:type"), "dxa")

def add_title(text, subtitle=None):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(8); p.paragraph_format.space_after = Pt(8)
    font(p.add_run(text), 29, True, NAVY)
    if subtitle:
        s = doc.add_paragraph(); s.paragraph_format.space_after = Pt(18)
        font(s.add_run(subtitle), 13.5, False, MID)

def add_bullets(items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet"); p.paragraph_format.space_after = Pt(4)
        font(p.add_run(item), 10.3, False, DARK)

def add_numbered(items):
    for item in items:
        p = doc.add_paragraph(style="List Number"); p.paragraph_format.space_after = Pt(5)
        font(p.add_run(item), 10.3, False, DARK)

def table(headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers)); t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False; t.style = "Table Grid"
    for i, h in enumerate(headers):
        c=t.rows[0].cells[i]; shade(c, NAVY); margins(c); c.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p=c.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER; font(p.add_run(h), 9.2, True, WHITE)
    header_props = t.rows[0]._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader"); repeat.set(qn("w:val"), "true"); header_props.append(repeat)
    for row in rows:
        cells=t.add_row().cells
        for i, value in enumerate(row):
            c=cells[i]; margins(c); c.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if len(t.rows)%2==1: shade(c, LIGHT)
            p=c.paragraphs[0]; font(p.add_run(str(value)), 9.0, False, DARK)
    if widths:
        for row in t.rows:
            for i,w in enumerate(widths): row.cells[i].width=Inches(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t

page_break_count = 0
def page_break():
    global page_break_count
    page_break_count += 1
    if page_break_count == 1:
        doc.add_page_break()

# Header/footer
hp=sec.header.paragraphs[0]; hp.alignment=WD_ALIGN_PARAGRAPH.RIGHT
font(hp.add_run("CAMPAIGN A–Z | 통합 선거 캠프 운영 플랫폼"), 8.5, True, MID)
fp=sec.footer.paragraphs[0]; fp.alignment=WD_ALIGN_PARAGRAPH.CENTER
font(fp.add_run("별도 신규사업 기획안 · 기존 개발과 독립 추진"), 8, False, MID)

# Cover
doc.add_paragraph().paragraph_format.space_after = Pt(80)
p=doc.add_paragraph(); font(p.add_run("PRODUCT PROPOSAL"), 11, True, CYAN)
add_title("선거의 A부터 Z까지", "AI가 캠페인 전략을 설계하고 조직별 마일스톤·일정·업무로 연결하는 통합 선거 캠프 운영 플랫폼")
p=doc.add_paragraph(style="Proposal Lead")
p.add_run("“이 서비스 하나만 구독하면 선거 캠프를 만들고 운영할 수 있다.”")
table(["문서 성격","제안 제품","추진 방식"],[["별도 신규사업 기획안","Campaign A–Z / 선거 캠프 OS","기존 개발 유지 + 선택적 모듈 연결"]],[1.7,2.4,2.4])
doc.add_paragraph("본 기획안은 현재 진행 중인 카드뉴스·디지털 명함 중심 개발을 중단하거나 변경하기 위한 문서가 아니다. 기획사가 신규 범위에 동의하는 경우에만 별도 프로젝트로 착수하고, 기존 결과물은 재사용 가능한 모듈로 연결한다.", style="Callout")
doc.add_paragraph("ZERO-VISIBILITY SECURITY — 캠프 데이터는 캠프만 열람하며, 운영사는 계약기간·요금제·서비스 상태 등 최소 운영정보만 확인한다.", style="Callout")
page_break()

add_title("1. 제안 개요")
doc.add_heading("1.1 제안 배경", 1)
doc.add_paragraph("선거 캠프에서는 전략 수립, 조직 운영, 후보 일정, 유세 현장, 메시지, 보도자료, 홍보물 제작, 임명장, 자료 보관이 서로 다른 도구와 메신저에 흩어진다. 이 때문에 최신 자료와 승인 상태를 찾기 어렵고, 일정 충돌·업무 누락·잘못된 메시지 사용·권한 과다 문제가 반복된다.")
doc.add_heading("1.2 제품 정의", 1)
doc.add_paragraph("Campaign A–Z는 디자인 도구의 묶음이 아니라 선거 전략을 실행 단위로 전환하고, 사람·일정·업무·콘텐츠·현장을 하나의 기록으로 연결하는 선거 캠프 운영 OS다.", style="Proposal Lead")
doc.add_heading("1.3 핵심 가치", 1)
table(["가치","제품이 제공하는 변화"],[
    ["한눈에 보는 캠프","후보·조직·콘텐츠·현장 상태를 오늘의 상황판에서 확인"],
    ["전략의 실행화","목표 → 전략 → 마일스톤 → 태스크 → 결과로 연결"],
    ["승인된 메시지","공약·수치·표현을 단일 원본으로 관리하고 제작물에 재사용"],
    ["역할별 업무","사용자는 자신에게 필요한 일정·업무·자료만 확인"],
    ["기록과 책임","변경·승인·발행·열람 이력을 감사 로그로 보존"]
],[1.65,4.85])

doc.add_heading("1.4 제품 모토", 1)
doc.add_paragraph("선거의 A부터 Z까지 — 캠프를 만들고, 운영하고, 기록하고, 홍보하는 모든 일을 하나의 서비스에서.", style="Callout")
page_break()

add_title("2. 사용자와 권한 체계")
doc.add_heading("2.1 주요 사용자", 1)
table(["역할","주요 업무","기본 접근 범위"],[
    ["후보자","일정 확인, 메시지·콘텐츠 최종 승인","전체 요약 및 승인 대상"],
    ["캠프장/상황실장","전략·조직·일정·업무 총괄","캠프 전체"],
    ["전략·정책팀","공약, 메시지, 근거자료, 전략 관리","전략·정책 공간"],
    ["홍보·공보팀","콘텐츠, SNS, 보도자료 제작·발행","홍보·미디어 공간"],
    ["조직팀","조직도, 참여자, 임명장, 지역 책임자","조직·신청자 정보"],
    ["유세팀","차량, 동선, 장비, 현장 일정","현장·유세 공간"],
    ["지역 책임자","지정 지역 일정·업무·결과 보고","배정 지역"],
    ["외부 제작자/자원봉사자","배정 업무 수행·자료 제출","배정 항목만"]
],[1.25,2.4,2.85])
doc.add_heading("2.2 권한 원칙", 1)
add_bullets([
    "역할 권한과 데이터 범위 권한을 분리한다.",
    "조회·작성·수정·승인·발행·삭제 권한을 각각 설정한다.",
    "개인정보, 위치정보, 전략문서, 회계자료는 별도의 보안 등급을 적용한다.",
    "조직 변경이나 퇴사 시 권한을 즉시 회수하고 접근 이력을 보존한다.",
    "중요 콘텐츠는 작성자와 승인자를 분리한다."
])
page_break()

add_title("3. 통합 대시보드와 전략 운영")
doc.add_heading("3.1 오늘의 캠프 상황판", 1)
add_bullets([
    "오늘 후보 일정과 이동 준비 상태",
    "유세차량 현재 위치 및 다음 일정",
    "긴급·지연·승인 대기 업무",
    "오늘 발행할 카드뉴스·메시지·보도자료",
    "이번 주 전략 마일스톤 진행률",
    "조직별 주요 보고와 위험 알림",
    "임명장 신청·승인·발송 현황"
])
doc.add_heading("3.2 전략 실행 구조", 1)
table(["단계","예시","시스템 연결"],[
    ["목표","청년층 인지도 상승","핵심 지표 설정"],
    ["전략","청년 주거 의제 선점","전략 담당·기간 지정"],
    ["마일스톤","청년 공약 발표","후보 일정·발표일 연결"],
    ["캠페인","청년 100인 인터뷰","콘텐츠·현장 프로젝트 생성"],
    ["태스크","섭외·촬영·보도자료·SNS","담당자·승인자·마감일"],
    ["결과","참여·보도·도달·반응","결과 보고와 자료 축적"]
],[1.1,2.25,3.15])
doc.add_heading("3.3 일정·태스크 관리", 1)
add_bullets(["일일·반복 업무", "마일스톤 역산 일정", "후보 일정과 캠프 일정", "부서별 칸반·캘린더", "지연·충돌 자동 경고", "회의용 일일·주간 보고서"])
doc.add_heading("3.4 AI 캠페인 전략 설계", 1)
doc.add_paragraph("AI 전략 엔진은 선거의 기본 조건과 캠프의 실제 역량을 입력받아 실행 가능한 캠페인 전략 초안을 작성한다. 단순 문서 생성에 그치지 않고 승인된 전략을 마일스톤·조직별 일정·담당자 태스크로 구조화한다.", style="Proposal Lead")
table(["AI 입력","AI 전략 산출물","운영 반영"],[
    ["선거 종류·지역·선거일","핵심 목표와 단계별 전략","선거일까지 역산한 마일스톤"],
    ["후보 정보·강점·핵심 의제","유권자·지역·이슈별 캠페인","전략별 책임 조직과 담당자"],
    ["조직도·인원·역할·가용시간","주간·일일 실행 과제","개인·조직 캘린더와 태스크"],
    ["예산·채널·차량·콘텐츠 역량","콘텐츠·현장·조직 실행안","콘텐츠 요청·유세 일정·보고"],
    ["승인된 공약·메시지·자료","위험·선행조건·대안 전략","지연·충돌·의존성 자동 경고"]
],[1.85,2.35,2.3])
doc.add_heading("3.5 전략 승인과 자동 전개", 1)
add_numbered([
    "캠프장이 선거 목표·후보·지역·기간·조직·예산을 입력한다.",
    "AI가 근거와 가정을 구분한 2~3개의 캠페인 전략안을 제안한다.",
    "전략팀이 목표·메시지·우선순위·위험을 검토하고 수정한다.",
    "후보 또는 승인권자가 최종 전략과 적용 기간을 승인한다.",
    "시스템이 승인 전략을 마일스톤, 조직별 일정, 담당자 태스크로 자동 변환한다.",
    "담당자는 자신의 화면에서 배정된 전략·일정·업무를 확인하고 실행 결과를 보고한다.",
    "결과 데이터가 누적되면 AI가 지연·위험·재조정안을 제안하고 재승인을 요청한다."
])
doc.add_heading("3.6 사람·조직별 전략 화면", 1)
table(["사용자","보이는 내용"],[
    ["후보자","전체 전략 요약, 핵심 마일스톤, 오늘 일정, 승인 대기, 주요 위험"],
    ["캠프장·상황실","전체 조직의 마일스톤, 일정 충돌, 지연 업무, 전략 진행률"],
    ["팀장","자기 조직의 전략 목표, 팀 일정, 담당자별 태스크, 보고·승인 항목"],
    ["실무자","본인에게 배정된 마일스톤, 오늘·이번 주 일정, 태스크, 필요한 자료"],
    ["지역 책임자","지정 지역 전략, 현장 일정, 조직 활동, 지역별 결과와 알림"],
    ["외부 제작자·봉사자","배정된 일정·제작 요청·제출 기한만 표시"]
],[1.8,4.7])
page_break()

add_title("4. 기능 모듈")
modules = [
    ("AI 캠페인 전략실","선거 조건과 후보·조직·기간·예산을 분석해 전략안을 제시하고, 승인 전략을 마일스톤·일정·태스크로 자동 전개한다."),
    ("콘텐츠 스튜디오","카드뉴스, 디지털 명함, 임명장, 현수막, 피켓, 공약서, SNS 이미지, 영상 자막, 인쇄 홍보물을 제작한다."),
    ("메시지 센터","핵심 메시지, 슬로건, 공약별 표준 표현, 연설문, 예상 문답, 위기 대응 문안을 승인 상태와 함께 관리한다."),
    ("보도자료·지식창고","보도자료, 통계, 공약 근거, 언론 기사, 회의록, 사진·영상 원본을 검색·버전·출처 기준으로 관리한다."),
    ("조직·임명장","신청, 개인정보 동의, 승인, 직책 배정, 임명장 생성·발송, QR 검증, 해촉까지 처리한다."),
    ("후보·캠프 일정","후보 일정, 수행 인력, 장소, 이동시간, 준비물, 공개 범위, 변경 알림을 통합한다."),
    ("유세·차량·동선","실시간 위치, 운행 일정, 이동시간, 중복 배정, 준비시간, 담당자, 현장 체크인을 관리한다."),
    ("알림·협업","조직별 공지, 태스크 댓글, 멘션, 승인 요청, 읽음 확인, 모바일 푸시를 제공한다."),
    ("성과·감사","전략별 실행률, 콘텐츠 발행, 현장 결과, 변경·승인·열람 기록을 통합한다.")
]
for title, body in modules:
    doc.add_heading(title, 2); doc.add_paragraph(body)
page_break()

add_title("5. 핵심 업무 흐름")
doc.add_heading("5.1 콘텐츠 제작·발행", 1)
add_numbered(["전략 또는 일정에서 콘텐츠 요청 생성","승인된 메시지·공약·자료 불러오기","카드뉴스·홍보물 초안 제작","팩트·표현·법률 검토","후보 또는 승인권자 결재","채널별 발행 및 링크 기록","성과와 최종 파일을 지식창고에 보존"])
doc.add_heading("5.2 임명장", 1)
add_numbered(["신청서 제출 및 목적별 동의","중복·조직·직책 검토","승인 또는 보완 요청","임명장 자동 생성","문자·이메일 발송","QR 검증 및 조직도 반영","해촉·만료·재발급 기록"])
doc.add_heading("5.3 유세 일정과 동선", 1)
add_numbered(["후보·유세차·행사 일정 입력","이동 및 준비시간 자동 계산","중복 차량·기사·장비 점검","위험·주의 장소와 경로 확인","운행 시작 후 위치 공유","도착·출발 체크","현장 결과와 사진 보고"])
page_break()

add_title("6. 기존 개발 소스 연결 전략")
doc.add_paragraph("현재 개발은 그대로 진행하고, 신규 기획은 별도 저장소·별도 계약·별도 일정으로 운영한다. 채택 시 기존 기능을 복사해 다시 만드는 것이 아니라 API와 공용 모듈로 연결한다.", style="Proposal Lead")
table(["구분","현재 자산","신규 기획에서의 활용"],[
    ["즉시 재사용","카드뉴스 캔버스 편집기","콘텐츠 스튜디오 모듈로 연결"],
    ["즉시 재사용","디지털 명함·모바일 공개 페이지","후보·조직원 프로필 모듈로 연결"],
    ["재사용 가능","임명장 생성·신청 흐름","조직 승인·발송·검증 기능으로 확장"],
    ["공통화 필요","회원·프로젝트·파일 저장","캠프·조직 단위 멀티테넌트 구조로 확장"],
    ["신규 개발","전략·마일스톤·태스크","캠프 운영 핵심 도메인"],
    ["신규 개발","일정·동선·실시간 위치","지도·위치 API 연동"],
    ["신규 개발","메시지·지식창고·승인","검색·버전·감사 체계"],
    ["외부 연동 우선","문자·이메일·지도·푸시","전문 사업자 API 사용"]
],[1.2,2.2,3.1])
doc.add_heading("6.1 분리 원칙", 1)
add_bullets([
    "기존 서비스의 화면·DB를 신규 기획 때문에 임의 변경하지 않는다.",
    "공용 기능은 독립 패키지 또는 API로 분리해 양쪽에서 사용한다.",
    "기존 데이터를 신규 서비스로 옮길 때는 명시적 마이그레이션을 사용한다.",
    "브랜드·요금제·계약 범위·배포 환경을 각각 분리할 수 있게 설계한다.",
    "기획사가 신규 범위를 채택하지 않아도 현재 납품물은 정상 완성된다."
])
page_break()

add_title("7. 기술·데이터 구조")
doc.add_heading("7.1 권장 아키텍처", 1)
add_bullets([
    "웹·모바일 반응형 프론트엔드",
    "캠프별 데이터가 분리되는 멀티테넌트 API",
    "역할 기반 권한(RBAC) + 지역·프로젝트 범위 권한",
    "파일·이미지·영상 객체 저장소와 버전 관리",
    "일정·업무·메시지·콘텐츠 이벤트 알림 서비스",
    "지도·경로·위치, 문자·이메일, 푸시 외부 API",
    "검색 가능한 지식창고와 승인된 자료 기반 AI 보조",
    "전략 초안·가정·근거·승인 버전을 분리하는 AI 전략 오케스트레이션",
    "승인 전략을 마일스톤·일정·태스크로 변환하는 워크플로 엔진",
    "조직·역할·지역·담당자에 따라 전략 화면을 구성하는 개인화 서비스",
    "감사 로그·백업·보존기간·자동 파기"
])
doc.add_heading("7.2 주요 데이터 객체", 1)
table(["영역","핵심 객체"],[
    ["캠프","캠프, 선거, 조직, 팀, 사용자, 역할, 권한"],
    ["전략","목표, 전략, 지표, 마일스톤, 캠페인"],
    ["운영","일정, 태스크, 체크리스트, 보고, 알림"],
    ["콘텐츠","제작물, 템플릿, 메시지, 승인, 발행, 성과"],
    ["지식","문서, 출처, 버전, 태그, 관련 공약"],
    ["현장","행사, 차량, 기사, 장비, 경로, 위치, 체크인"],
    ["조직","신청자, 직책, 임명장, 발급, 해촉"]
],[1.2,5.3])
page_break()

add_title("8. 보안·법률·운영 기준")
doc.add_paragraph("보안 원칙: 캠프 데이터는 캠프만 볼 수 있다. 플랫폼 운영회사와 구독서비스 회사도 고객의 콘텐츠·전략·조직·일정·메시지·위치정보를 열람할 수 없도록 캠프별 보안 컨테이너로 격리한다.", style="Proposal Lead")
doc.add_paragraph("정치적 견해·정당 관련 정보, 연락처, 위치정보를 다루므로 보안과 준법은 부가 기능이 아니라 제품의 기본 구조다. 운영사는 서비스 계약 유지에 필요한 최소 메타데이터만 처리하고 캠프 업무 데이터의 평문에는 기본적으로 접근하지 못한다.")
doc.add_heading("8.1 캠프별 보안 컨테이너", 1)
table(["보안 원칙","적용 방식"],[
    ["테넌트 완전 격리","캠프별 데이터베이스 논리 격리, 파일 저장 경로 분리, 교차 조회 차단"],
    ["고객 데이터 암호화","전송·저장 암호화와 캠프별 암호키 분리. 가능하면 고객 관리 키 또는 외부 키 관리 적용"],
    ["운영자 기본 접근 차단","관리자 화면에서도 캠프 콘텐츠·문서·메시지·위치·개인정보 원문을 표시하지 않음"],
    ["최소 메타데이터 원칙","구독사는 계약 시작일·종료일, 요금제, 계정 수, 저장량, 서비스 상태만 확인"],
    ["백업도 동일 보호","백업·로그·검색색인·캐시에도 동일한 격리와 암호화 정책 적용"],
    ["캠프 종료 통제","계약 종료 시 고객 선택에 따라 데이터 반출 후 자동 잠금·보존·완전 파기"]
],[1.65,4.85])
doc.add_heading("8.2 운영사가 확인할 수 있는 정보", 1)
table(["확인 가능","확인 불가"],[
    ["캠프 식별용 계약번호","전략·공약·메시지 내용"],
    ["계약 시작일·종료일","후보·조직원 일정과 태스크"],
    ["요금제·라이선스 수","보도자료·회의록·첨부문서 원문"],
    ["저장용량·API 사용량","연락처·신청자·임명장 개인정보"],
    ["서비스 정상·장애 상태","차량 위치·후보 동선·현장 보고"],
    ["청구·결제 상태","카드뉴스·홍보물의 미공개 초안"]
],[3.25,3.25])
doc.add_heading("8.3 예외 접근과 투명성", 1)
doc.add_paragraph("현실적인 SaaS 운영에서는 장애 복구, 고객이 요청한 기술지원, 법적 의무 때문에 제한적 접근 절차가 필요할 수 있다. 이 경우에도 운영자가 임의로 열람할 수 없고 다음의 ‘봉인 해제’ 절차를 거쳐야 한다.", style="Callout")
add_bullets([
    "고객 캠프의 명시적 요청 또는 유효한 법적 근거가 있을 때만 접근한다.",
    "고객 승인자와 보안책임자 등 복수 승인을 요구한다.",
    "접근 대상·목적·담당자·시작·종료 시간을 사전에 한정한다.",
    "모든 화면 조회·다운로드·변경을 위변조 방지 감사 로그로 남긴다.",
    "접근 종료 즉시 임시 권한과 복호화 세션을 폐기하고 고객에게 결과를 통지한다.",
    "가능한 운영 절차는 콘텐츠 원문을 보지 않는 진단 정보와 익명화 로그로 처리한다."
])
doc.add_heading("8.4 개인정보·선거 준법", 1)
add_bullets([
    "개인정보와 민감정보 동의를 분리하고 수집 목적·출처·보유기간을 기록한다.",
    "위치 공유는 근무시간에만 사용하고 당사자가 시작·중단 상태를 확인할 수 있게 한다.",
    "문자·이메일은 발송 권한, 수신 거부, 발송 이력, 적용 선거 규칙을 확인한다.",
    "허위사실·비방·출처 누락 방지를 위한 검토 체크리스트와 승인 단계를 둔다.",
    "중요 자료 암호화, 다중인증, 세션 제한, 다운로드 통제를 적용한다.",
    "캠프 종료 시 데이터 이관·보관·파기 절차를 제공한다.",
    "시스템은 법률 적합성을 자동 보증하지 않고 법률 담당자의 검토와 선관위 질의 결과를 기록한다."
])
doc.add_heading("8.5 필수 감사 기록", 1)
table(["기록 대상","남겨야 할 정보"],[
    ["콘텐츠","작성·수정·검토·승인·발행자와 시각"],
    ["개인정보","조회·다운로드·수정·삭제 이력"],
    ["권한","부여·변경·회수 담당자와 사유"],
    ["메시지","승인 버전과 실제 사용 제작물"],
    ["위치","공유 동의·시작·종료·조회 이력"]
],[1.55,4.95])
page_break()

add_title("9. 단계별 개발 로드맵")
table(["단계","범위","완료 기준"],[
    ["Phase 0\n기획·검증","캠프 인터뷰, 권한·업무 흐름, 법률 검토, 정보구조","PRD·프로토타입·데이터 정책 승인"],
    ["Phase 1\n운영 MVP","캠프/조직, 권한, AI 전략 초안, 마일스톤, 일정, 태스크, 자료실, 개인화 대시보드","실제 캠프 1곳에서 전략 수립과 일일 운영 가능"],
    ["Phase 2\n제작 통합","카드뉴스, 명함, 임명장, 메시지, 승인, 발행 기록","기존 제작 모듈과 운영 업무 연결"],
    ["Phase 3\n현장 운영","유세 일정, 차량, 경로, 위치, 체크인, 모바일 알림","후보·차량 동선 충돌 사전 검증"],
    ["Phase 4\n미디어·분석","보도자료, 채널 캘린더, 발송 연동, 성과·AI 보조","전략-실행-성과 통합 보고"]
],[1.15,2.45,2.9])
doc.add_heading("9.1 우선순위 원칙", 1)
add_bullets([
    "채팅보다 일정·태스크·승인·대시보드를 먼저 만든다.",
    "자체 문자·지도 엔진보다 검증된 외부 API를 우선 연결한다.",
    "AI 생성보다 승인된 자료와 변경 이력의 신뢰성을 먼저 확보한다.",
    "AI 전략은 제안으로만 생성하며 사람의 승인 전에는 일정·마일스톤에 적용하지 않는다.",
    "모든 기능을 한 번에 만들지 않고 실제 캠프 1곳의 운영 흐름으로 검증한다."
])
page_break()

add_title("10. 사업·계약 제안")
doc.add_heading("10.1 제안 방식", 1)
table(["선택안","설명","영향"],[
    ["A. 기존 개발만 진행","현재 카드뉴스·디지털 명함 등 약정 범위 완성","기존 일정과 계약 유지"],
    ["B. 별도 기획 검토","본 문서를 기준으로 조사·설계만 별도 수행","개발 착수 전 범위·비용 확정"],
    ["C. 신규 제품 착수","Campaign A–Z를 별도 프로젝트로 개발","기존 소스는 합의된 모듈만 재사용"],
    ["D. 단계적 통합","기존 납품 후 운영 MVP부터 연결","위험과 초기 비용을 분산"]
],[1.45,2.65,2.4])
doc.add_heading("10.2 계약서에 명확히 할 항목", 1)
add_bullets([
    "기존 소스의 소유권·사용권·2차 개발 권한",
    "공용 모듈의 유지보수 책임과 버전 정책",
    "신규 기획의 별도 범위·일정·검수 기준",
    "외부 API 비용과 운영비 부담 주체",
    "개인정보처리자·수탁자 역할과 사고 대응",
    "기획사가 신규 사업을 채택하지 않을 때 기존 납품 범위"
])
doc.add_heading("10.3 의사결정 요청", 1)
doc.add_paragraph("우선 본 기획안을 신규 사업 후보로 검토하고, 기획사와 함께 ① 사업 채택 여부, ② 기존 소스 재사용 권한, ③ 1단계 MVP 범위, ④ 실제 캠프 파일럿 가능성을 결정한다.", style="Callout")
doc.add_page_break()

add_title("11. 결론")
doc.add_paragraph("Campaign A–Z는 선거 홍보물 제작 서비스를 넘어, 전략이 업무가 되고 업무 결과가 콘텐츠·보도자료·현장 기록으로 이어지는 통합 선거 캠프 운영 플랫폼이다.", style="Proposal Lead")
doc.add_paragraph("기술적으로 개발 가능하며 현재 개발된 카드뉴스 편집기·디지털 명함·임명장 기능을 재사용할 수 있다. 다만 신규 기획 채택 여부가 정해지지 않은 현재는 기존 개발을 그대로 완성하고, 본 기획을 별도 제품·별도 계약·별도 저장소로 관리하는 것이 가장 안전하다.")
doc.add_paragraph("권장 결론: 기존 개발은 현 범위대로 진행하고, Campaign A–Z는 Phase 0 기획·검증을 별도 발주한 뒤 운영 MVP 착수 여부를 결정한다.", style="Callout")

doc.core_properties.title = "선거의 A부터 Z까지 - 통합 선거 캠프 운영 플랫폼 기획안"
doc.core_properties.subject = "기존 개발과 분리된 신규 사업 제안"
doc.core_properties.author = "POLLINSIGHT"
doc.save(OUT)
print(OUT)
